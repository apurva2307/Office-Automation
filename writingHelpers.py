from puToFullNameMap import puMap, summaryMap
from dataHelpers import *
from docx.enum.text import WD_ALIGN_PARAGRAPH

puNameMap = puMap()
summaryNameMap = summaryMap()


def moreLess(data):
    return "more than" if data > 0 else "less than"


def iteratePara(p, data):
    for index, pu in enumerate(data.keys()):
        if index == len(data) - 1:
            p.add_run(f"and {puNameMap[pu]} ({data[pu]}%).")
        elif index == len(data) - 2:
            p.add_run(f"{puNameMap[pu]} ({data[pu]}%) ")
        else:
            p.add_run(f"{puNameMap[pu]} ({data[pu]}%), ")


def iterateParaSumm(p, data):
    for index, value in enumerate(data.keys()):
        p.add_run(f"{summaryNameMap[value]} ({data[value]}%), ")

def sofWiseSlowMore(dataCapex, p, sof, budType, margin):
    slowProgCapexCap = slowProgCapex(dataCapex, sof, margin)
    highProgCapexCap = highProgCapex(dataCapex, sof, margin)
    budUtil = dataCapex["TOTAL"][sof]["NCR"][-1]
    p.add_run(f"Overall expenditure is {budUtil}% of {budType}. ")
    if len(slowProgCapexCap) > len(highProgCapexCap):
        p.add_run("Progress under all Plan Heads is slow except under ")
        for index, key in enumerate(highProgCapexCap.keys()):
            if index == len(highProgCapexCap)-1:
                p.add_run(f"and {key} where progress is more.")
                break
            p.add_run(f"{key}, ")
    else:
        p.add_run("Progress under all Plan Heads is high except under ")
        for index, key in enumerate(slowProgCapexCap.keys()):
            if index == len(slowProgCapexCap)-1:
                p.add_run(f"and {key} where progress is slow.")
                break
            p.add_run(f"{key}, ")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY