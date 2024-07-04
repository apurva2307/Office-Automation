from data import extractData, extractDataSummary, extractDataCapex
from docx import Document
from datetime import datetime
from puToFullNameMap import puMap
from dataHelpers import *
from writingHelpers import *
from docx.enum.text import WD_ALIGN_PARAGRAPH

FRMonth = "APR24"
filePath = f"files/OWE-{FRMonth}.xlsx"
month = "Jan' 22"
budType = "RG"
marginExcessBud = 5
marginExcessCoppy = 20
marginExLessCapex = 5

monthdata = extractData(filePath)
dataCapex = extractDataCapex("files/Capex Review 2021-22.xlsx", "Capex Jan-22")
currentMonth = datetime.now().month
frMonth = currentMonth - 4 if currentMonth > 4 else currentMonth + 8
puNameMap = puMap()
highUtilPuStaff = highUtilStaff(monthdata, marginExcessBud)
highUtilPuStaffCoppy = highUtilStaffCoppy(monthdata, marginExcessCoppy)
highUtilPuNonStaff = highUtilNonStaff(monthdata, marginExcessBud)
highUtilPuNonStaffCoppy = highUtilNonStaffCoppy(monthdata, marginExcessCoppy)
budget, netTotal, budgetUtil, varCPer = getMainData(monthdata, "NET")
staffBud, staffNet, staffUtil, staffVarC = getMainData(monthdata, "STAFF")
nonStaffBud, nonStaffNet, nonStaffUtil = (
    budget - staffBud,
    round((netTotal - staffNet), 2),
    round(((netTotal - staffNet) * 100 / (budget - staffBud)), 2),
)
nonStaffCoppy = round(monthdata["NET"]["toEndActualsCoppy"][-1] / 10000, 2) - round(
    monthdata["STAFF"]["toEndActualsCoppy"][-1] / 10000, 2
)
nonStaffVarC = round(((nonStaffNet - nonStaffCoppy) * 100 / nonStaffCoppy), 2)
document = Document()
document.add_heading(f"NCR Financial Review {FRMonth[:3]}-20{FRMonth[3:]}", level=1)
document.add_heading("Revenue Expenditure:", level=1)
document.add_paragraph(
    f"The Revised Grant for Ord. Working Expenses (OWE) 2021-22, excluding suspense is Rs {budget} crore, more than SL by Rs. 430.10 crore and more than last year actuals by only Rs. 889.01 crore (11.37%).",
    style="List Bullet",
)
document.add_paragraph(
    f"OWE (excluding suspense) to end {month} amounts to Rs {netTotal} crore which is {budgetUtil}% of the {budType}, and {moreLess(varCPer)} COPPY by {varCPer}%.",
    style="List Bullet",
)
p1 = document.add_paragraph(
    f"Staff expenditure to end {month} of {staffNet} crore is {staffUtil}% of {budType}, and {moreLess(staffVarC)} COPPY by {staffVarC}%. Utilisation of {budType} is high for ",
    style="List Bullet",
)
iteratePara(p1, highUtilPuStaff)
p1.add_run(" Growth over last year is high for ")
iteratePara(p1, highUtilPuStaffCoppy)
p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
dataOther = extractDataSummary(filePath)
highUtilPuNonStaffOther = highUtilNonStaffOther(dataOther, marginExcessBud)
highUtilPuNonStaffOtherCoppy = highUtilNonStaffOther(dataOther, marginExcessCoppy)
p2 = document.add_paragraph(
    f"Non-Staff expenditure to end {month} of {nonStaffNet} crore is {nonStaffUtil}% of {budType} and {moreLess(nonStaffVarC)} COPPY by {nonStaffVarC}%. Utilisation of {budType} is high for ",
    style="List Bullet",
)
iterateParaSumm(p2, highUtilPuNonStaffOther)
iteratePara(p2, highUtilPuNonStaff)
p2.add_run(" Growth over last year is high for ")
iterateParaSumm(p2, highUtilPuNonStaffOtherCoppy)
iteratePara(p2, highUtilPuNonStaffCoppy)
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_heading("CAPEX:", level=1)
grossCapex = dataCapex["G-TOTAL"]["NCR"]
p3 = document.add_paragraph(
    f"Revised Grant for current year 2021-22 for CAPEX (Gross excluding suspense) is Rs. {inCrore(grossCapex[0])} crore. Expenditure (Gross excluding suspense) to end of {month} is Rs. {inCrore(grossCapex[1])} crore, which is {grossCapex[2]} % of the {budType}.",
    style="List Bullet",
)
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
document.add_paragraph(
    "Progress of expenditure under various sources is as under:", style="List Bullet"
)
p4 = document.add_paragraph("", style="List Bullet 2")
p4.add_run("Capital (excluding Suspense):- ").bold = True
sofWiseSlowMore(dataCapex, p4, "CAP", budType, marginExLessCapex)
p5 = document.add_paragraph("", style="List Bullet 2")
p5.add_run("DF:- ").bold = True
sofWiseSlowMore(dataCapex, p5, "DF", budType, marginExLessCapex)
p6 = document.add_paragraph("", style="List Bullet 2")
p6.add_run("DRF:- ").bold = True
sofWiseSlowMore(dataCapex, p6, "DRF", budType, marginExLessCapex)
p7 = document.add_paragraph("", style="List Bullet 2")
p7.add_run("RRSK:- ").bold = True
sofWiseSlowMore(dataCapex, p7, "RRSK", budType, marginExLessCapex)
p8 = document.add_paragraph("", style="List Bullet 2")
p8.add_run("Operating Ratio: ").bold = True
p8.add_run(
    f"Adding the {budType} for OWE of Rs. {budget} crores, appropriation to DRF of Rs. 20 crore and Pension fund of Rs. 2612 crore, the target for Gross expenditure (without suspense) for 2021-22 is Rs. 11339.05 crore. With target Gross revenue of Rs. 15470.62 crore, the target for Operating ratio for the year is 73.29%. "
)
p8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p9 = document.add_paragraph(
    "Operating Ratio to end Jan’22 is 85.66%, more than the target Operating Ratio but less than Operating Ratio of 94% to end Jan’21, when revenues were down due to Covid lockdown."
)
p9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

document.save(f"FR_{FRMonth}.docx")
